"""
Generate the LegalLens technical report as a Word document.

Kept in the repo so the report can be regenerated after code changes rather than
hand-edited and drifting out of date:

    python docs/build_report.py
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent.parent / "LegalLens_Technical_Report.docx"

INK    = RGBColor(0x1A, 0x1A, 0x1A)
MUTED  = RGBColor(0x5F, 0x6B, 0x7A)
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
CODE_BG = "F4F5F7"
HEAD_BG = "1F4E79"


# ── Formatting helpers ────────────────────────────────────────────────────────

def shade(cell, hex_fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def code(doc, text):
    """Monospaced, shaded block for code and configuration."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1F, 0x33, 0x46)
    pPr = p._p.get_or_add_pPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), CODE_BG)
    pPr.append(el)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cell, HEAD_BG)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(value))
            run.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def body(doc, text, italic=False, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.font.color.rgb = MUTED if italic else INK
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            lead, rest = item
            r = p.add_run(lead)
            r.bold = True
            r.font.size = Pt(10.5)
            r2 = p.add_run(rest)
            r2.font.size = Pt(10.5)
        else:
            p.add_run(item).font.size = Pt(10.5)


def h(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = ACCENT
    return heading


def callout(doc, label, text):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    shade(cell, "FFF4E5")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}  ")
    r.bold = True
    r.font.size = Pt(9.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(9.5)
    doc.add_paragraph()


# ── Document ──────────────────────────────────────────────────────────────────

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# ── Title page ────────────────────────────────────────────────────────────────

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("LegalLens")
run.font.size = Pt(38)
run.bold = True
run.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Hybrid Semantic + Keyword Search for Legal Documents")
run.font.size = Pt(14)
run.font.color.rgb = MUTED

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run("Technical Report — Implementation, Data, Tooling, Evaluation and Output")
run.font.size = Pt(11)
run.italic = True
run.font.color.rgb = MUTED

doc.add_paragraph()

table(doc, ["Property", "Value"], [
    ("Report date", "25 July 2026"),
    ("Repository", "LegalLens (git, branch main)"),
    ("Language / runtime", "Python 3.11+"),
    ("Application code", "~3,400 lines across 24 modules"),
    ("Test suite", "82 tests, all passing"),
    ("Search backend", "Elasticsearch 8.14.0 (single node, Docker)"),
    ("Embedding model", "sentence-transformers/all-mpnet-base-v2 (768-dim)"),
    ("LLM", "claude-sonnet-4-6 (query understanding + answer synthesis)"),
    ("Interfaces", "Streamlit web app; interactive CLI"),
    ("Corpus indexed", "956 chunks — 10 PDFs, 5 emails, 823 LEDGAR provisions"),
], widths=[2.0, 4.7])

doc.add_page_break()

# ── 1. Executive summary ──────────────────────────────────────────────────────

h(doc, "1.  Executive Summary", 1)

body(doc,
     "LegalLens is a document search system for legal corpora. It answers natural-language "
     "questions over a mixed corpus of court opinions, contracts, email threads and regulatory "
     "provisions, and returns a written answer in which every claim is cited back to the source "
     "document it came from.")

body(doc,
     "The central design problem is that legal search has two irreconcilable requirements. A "
     "lawyer asking about \"circumstances preventing contract performance\" needs a system that "
     "understands this means force majeure — that is semantic search. The same lawyer asking "
     "about \"3 Allen Center\" or a specific citation number needs exact lexical matching, where "
     "an embedding model is actively unhelpful. LegalLens runs both retrievers on every query and "
     "fuses their rankings, so neither requirement is sacrificed for the other.")

h(doc, "1.1  Pipeline at a glance", 2)

code(doc,
     "Documents                 Ingestion            Chunking             Index\n"
     "─────────────────────────────────────────────────────────────────────────────────\n"
     "PDF (native + scanned) →  Docling OCR      →  RSM section tree  →  ┐\n"
     "Email (Enron format)   →  custom parser    →  one email = chunk →  ├→  Elasticsearch\n"
     "JSONL (LEDGAR)         →  ijson streaming  →  greedy KV merge   →  ┘    (single index,\n"
     "                                                  ↓                      768-dim vectors)\n"
     "                                          all-mpnet-base-v2\n"
     "                                          768-dim embeddings\n"
     "\n"
     "Query                     Retrieval            Presentation\n"
     "─────────────────────────────────────────────────────────────────────────────────\n"
     "natural language       →  Claude NLI       →  BM25  ┐\n"
     "                          (intent+filters)     kNN  ┴→ RRF → gate → group → confidence\n"
     "                                                              ↓\n"
     "                                              Claude synthesis → cited answer + sources")

h(doc, "1.2  What the system does end to end", 2)

bullets(doc, [
    ("Understands the question. ", "Claude parses the natural-language query into a retrieval "
     "string, an intent classification, and metadata filters (document type, date range, party names)."),
    ("Retrieves twice. ", "BM25 full-text and kNN vector search each return a ranked window of "
     "50 candidates from a single Elasticsearch index."),
    ("Fuses the rankings. ", "Reciprocal Rank Fusion merges the two lists in Python, with a "
     "relevance gate that returns nothing rather than returning noise."),
    ("Collapses duplicates. ", "Chunks from the same source document are merged into one cited "
     "source, so a 24-page opinion is cited once rather than three times."),
    ("Scores confidence. ", "A readable 0–100% confidence is reconstructed from cosine similarity "
     "and BM25 rank, and shown as a coloured ring beside each citation."),
    ("Writes a cited answer. ", "Claude synthesises a short answer from the retrieved excerpts, "
     "citing sources by number and declining to answer when the excerpts are insufficient."),
    ("Shows the evidence. ", "Every matched page of a source PDF is rendered with the matched "
     "passage highlighted in place."),
])

doc.add_page_break()

# ── 2. Technology stack ───────────────────────────────────────────────────────

h(doc, "2.  Tools and Technology Choices", 1)

body(doc,
     "Every dependency below was chosen against named alternatives. The rationale is recorded in "
     "DECISIONS.md in the repository; the summary is reproduced here.")

table(doc, ["Layer", "Choice", "Why this over the alternatives"], [
    ("OCR / parsing", "Docling (IBM)",
     "Runs fully locally with no API key; one pipeline handles both scanned and native PDFs; "
     "emits clean Markdown including tables. Rejected: pytesseract (poor on low-quality scans, "
     "weak tables), pdfplumber (no OCR at all), Gemini Vision and Mistral OCR (per-page cost and "
     "an external dependency that breaks local reproducibility)."),
    ("Search backend", "Elasticsearch 8.14",
     "The only candidate offering native BM25 and kNN in one store. Rejected: ChromaDB (no "
     "full-text search — fatal for clause numbers and party names), ChromaDB + rank_bm25 (two "
     "stores to keep in sync), OpenSearch (no advantage at prototype scale)."),
    ("Embeddings", "all-mpnet-base-v2",
     "768-dim, higher accuracy than MiniLM's 384-dim. Embedding cost is paid once at ingestion, "
     "so the ~3× slower encode is irrelevant to query latency. Rejected: OpenAI embeddings (per-chunk "
     "cost, external dependency), MiniLM (accuracy matters more than speed for legal recall)."),
    ("Large JSON", "ijson",
     "Streams records at constant memory. json.load() on a 200 MB compliance export needs 1–2 GB "
     "of RAM."),
    ("LLM", "claude-sonnet-4-6",
     "Two roles: query understanding and answer synthesis. System prompts are marked with ephemeral "
     "cache_control so the fixed instructions are cached across calls."),
    ("Web UI", "Streamlit 1.54",
     "Fastest path from Python functions to a working interface; no separate frontend build."),
    ("PDF rendering", "PyMuPDF (fitz)",
     "Renders pages to PNG and applies word-level highlight annotations for the in-app viewer."),
    ("Tests", "pytest",
     "82 tests split into fast unit tests (mocked ES) and integration tests (live ES), separated "
     "by marker."),
], widths=[1.1, 1.3, 4.3])

h(doc, "2.1  Deployment", 2)

body(doc,
     "docker-compose brings up Elasticsearch (single node, security disabled, JVM heap capped at "
     "512 MB) and the application container, which waits on an Elasticsearch health check before "
     "starting. The application image installs libgl1, libglib2.0-0 and libgomp1, which Docling "
     "requires for OCR.")

code(doc,
     "docker compose up -d elasticsearch     # start the search backend\n"
     "python ingest.py --all                 # ingest, chunk, embed, index everything\n"
     "streamlit run app.py                   # web interface at localhost:8501\n"
     "python search_cli.py                   # terminal interface\n"
     "python -m pytest tests -q              # 82 tests\n"
     "python -m pytest tests -q -m \"not integration\"   # unit tests only, no ES needed")

doc.add_page_break()

# ── 3. Data ───────────────────────────────────────────────────────────────────

h(doc, "3.  The Data", 1)

body(doc,
     "Three corpora were chosen because they demand three genuinely different ingestion and "
     "chunking strategies — OCR of visual documents, metadata-rich plain text, and "
     "schema-consistent structured records. This exercises more of the architecture than three "
     "similar formats would.")

h(doc, "3.1  Corpus inventory", 2)

table(doc, ["Corpus", "Format", "Volume", "Source"], [
    ("Court opinions and contracts", "PDF (native)", "5 documents, 90 chunks",
     "Public court opinions, 2026"),
    ("Court opinions and reports", "PDF (scanned, OCR)", "5 documents, 36 chunks",
     "Public records, 1976–2006"),
    ("Email threads", "Plain text", "5 messages, 7 chunks", "Enron corpus"),
    ("Regulatory provisions", "JSONL", "374,639 chunks available; 823 indexed",
     "LEDGAR (SEC filings, 2016–2019)"),
], widths=[2.0, 1.4, 1.6, 1.7])

h(doc, "3.2  Native PDFs", 2)

table(doc, ["Document", "Court / Type", "Chunks", "Subject"], [
    ("20260410_c378748_58_378748.opn.pdf", "Michigan Court of Appeals", "24",
     "Constitutional challenge to MCL 388.1631aa, conditioning school mental-health funds"),
    ("alterna_aircraft_v_b_ltd._v._spicejet_ltd..pdf", "Washington Supreme Court (En Banc)", "23",
     "Whether foreign-judgment recognition requires property in the forum state"),
    ("national_trust..._v._nps.pdf", "D.C. Circuit", "17",
     "Stay of injunction over White House East Wing ballroom demolition"),
    ("amin_v._taylor__francis_group_llc.pdf", "D.D.C.", "15",
     "Defamation claim over ICE detention-centre allegations"),
    ("arotin_v._arotin.pdf", "Ohio 11th Appellate District", "11",
     "Summary judgment on standing, statute of frauds, quitclaim deed"),
], widths=[1.9, 1.5, 0.6, 2.7])

h(doc, "3.3  Scanned PDFs (OCR path)", 2)

table(doc, ["Document", "Chunks", "Subject"], [
    ("cia-rdp81-00706r000100230029-5.pdf", "21", "Declassified CIA memo on the Anders Collection"),
    ("056.pdf", "9", "D.N.J. — Newark Coalition for Low Income Housing v. NRHA and HUD"),
    ("RJHC010111202006.pdf", "3", "Rajasthan High Court — motor accident compensation appeal"),
    ("19760202-backgrounder...miranda-kern.pdf", "2", "VOA radio backgrounder on Ernesto Miranda"),
    ("19780425-backgrounder...womens_pensions.pdf", "1", "VOA backgrounder on the pension-contribution ruling"),
], widths=[2.6, 0.7, 3.4])

body(doc,
     "These five exist specifically to exercise the OCR path and its failure modes: transcription "
     "artefacts ('THE CONSTITUTIONAL LAV!'), missing text layers, and non-US document conventions.",
     italic=True)

h(doc, "3.4  Emails", 2)

table(doc, ["File", "Date", "Subject", "Notes"], [
    ("email_1.txt", "2000-09-25", "Energy Bar Association Program", "Forwarded external message"),
    ("email_2.txt", "2001-10-25", "Weekend Outage Report", "Systems notice, no real sender"),
    ("email_3.txt", "2000-12-27", "RE: EWEB Schedule to the Master Agreement", "Contract negotiation thread"),
    ("email_4.txt", "2002-01-28", "RE: Super Bowl Party", "Off-topic — negative-control material"),
    ("email_5.txt", "2000-12-27", "Technical Questions", "Multi-recipient, long body → split into fragments"),
], widths=[1.0, 0.9, 2.4, 2.4])

body(doc,
     "The Enron export is not valid RFC 2822 — To: values are Python list reprs, File-Name is a "
     "non-standard header, and quoted history is inline rather than MIME. Python's email stdlib "
     "parser fails silently on all three, so a line-by-line parser was written instead.")

h(doc, "3.5  LEDGAR compliance provisions", 2)

body(doc,
     "LEDGAR contains contractual provisions extracted from SEC filings, each labelled with its "
     "clause type. Records carry three keys: provision (the clause text), label (a list of clause "
     "types) and source (the originating filing path).")

table(doc, ["Most frequent labels (first 60,000 records)", "Count"], [
    ("governing laws", "1,419"), ("counterparts", "999"), ("amendments", "916"),
    ("severability", "914"), ("entire agreements", "839"), ("notices", "737"),
    ("waivers", "698"), ("successors", "660"), ("expenses", "599"), ("assigns", "478"),
], widths=[4.6, 1.0])

body(doc, "8,871 distinct labels appear across the sampled records.", italic=True)

doc.add_page_break()

# ── 4. Implementation ─────────────────────────────────────────────────────────

h(doc, "4.  Implementation", 1)

h(doc, "4.1  Ingestion", 2)

body(doc, "PDF — ingestion/pdf_ingester.py")
bullets(doc, [
    ("Staleness detection. ", "SHA256 of the raw bytes is written to a .hash sidecar next to a "
     ".md cache. On re-ingestion an unchanged hash skips OCR entirely and reuses the cached "
     "Markdown. This directly addresses stale results: a document replaced under the same filename "
     "is detected and re-indexed, which an exists-check on the .md file would miss."),
    ("Page provenance. ", "Docling's item iterator is walked manually to inject <!-- page N --> "
     "markers at page boundaries, because the standard Markdown export discards page numbers. "
     "Section headers are also emitted manually, since per-item export does not add ## prefixes."),
    ("Date extraction. ", "Filename regex (YYYY-MM-DD) first, then four date patterns against the "
     "first page only."),
    ("Concurrency. ", "ThreadPoolExecutor with 3 workers and staggered submission."),
])

body(doc, "Email — ingestion/email_ingester.py")
bullets(doc, [
    ("Multi-line headers. ", "Cc and Bcc values spanning several lines are reattached by tracking "
     "the last header key and appending whitespace-prefixed continuation lines."),
    ("Reply chains. ", "The body is split on -----Original Message----- and ----- Forwarded by ----- "
     "markers, with separate extractors for the two different header layouts."),
    ("Threading. ", "thread_id is derived from the file path structure so related messages link "
     "without being concatenated, preserving individual message provenance."),
])

body(doc, "JSON — chunking/json_chunker.py")
bullets(doc, [
    ("Streaming. ", "ijson.items() with multiple_values=True reads the JSONL one record at a time "
     "at constant memory."),
    ("Citation metadata. ", "Each record's source filing and clause labels are captured onto the "
     "chunk as source_documents and provision_labels, so provisions cite their originating SEC "
     "filing rather than appearing as an anonymous fragment."),
])

h(doc, "4.2  Chunking — three strategies, one token budget", 2)

body(doc,
     "All three chunkers target a 512-token budget, estimated as words / 0.75. A universal "
     "character splitter was rejected: splitting a payment schedule mid-table or an email "
     "mid-thread destroys exactly the semantic coherence that retrieval depends on.")

table(doc, ["Type", "Unit", "Strategy"], [
    ("PDF text", "Section", "RSM: recursive split down a section → paragraph → sentence tree, then "
     "greedy merge of adjacent leaves sharing a heading context, up to 512 tokens."),
    ("PDF table", "Whole table", "Never split, at any size. A 51-row defendant list broken in half "
     "is useless — a name without its case and side has no retrieval value."),
    ("Email", "Message", "One email is one chunk, including quoted history so the embedding sees "
     "full context. Split at paragraph boundaries only if the body exceeds budget; fragments carry "
     "parent_chunk_index and fragment_index."),
    ("JSON", "Provision", "Whole record if it fits, greedily merged with neighbours. Oversized "
     "records split at key boundaries; a single oversized value falls back to a word split."),
], widths=[0.9, 1.0, 4.8])

body(doc, "Three post-processing passes clean up the PDF output:")
bullets(doc, [
    "Small headerless chunks following a table are appended to it — these are footnotes.",
    "Small chunks preceding a table are prepended to it — these are introductory labels.",
    "Consecutive section chunks are greedily merged while they fit the budget.",
])

h(doc, "4.3  Embedding", 2)

bullets(doc, [
    ("One function, all types. ", "embed_chunks() accepts PDF, email and JSON chunks alike — they "
     "share a .text field."),
    ("Lazy singleton. ", "The model is loaded once on first use and reused."),
    ("Normalised vectors. ", "normalize_embeddings=True produces unit vectors, so a dot product "
     "equals cosine similarity — faster in Elasticsearch and, as section 4.6 shows, the key to "
     "recovering an interpretable confidence."),
    ("Batching. ", "Batch size 64. Sentinel fields set to −1 are stripped before indexing."),
])

h(doc, "4.4  Storage and index design", 2)

body(doc,
     "A single Elasticsearch index, legallens, holds all three document types under one flat "
     "mapping. Per-type indices were rejected because cross-index queries complicate rank fusion "
     "for no schema benefit at this scale. Dynamic mapping was rejected because it mis-types "
     "fields — page_number infers as text — and dense vectors require explicit dims and similarity.")

code(doc,
     '"text":      {"type": "text"}                      # BM25 full-text\n'
     '"embedding": {"type": "dense_vector", "dims": 768,\n'
     '              "index": true, "similarity": "dot_product"}\n'
     '"chunk_type", "doc_type", "filename",              # exact-match keyword fields\n'
     '"sender", "date", "document_date", "content_hash",\n'
     '"source_documents", "provision_labels"\n'
     '"page_number", "token_count", "chunk_index"        # integer fields')

bullets(doc, [
    ("Document _id. ", "{filename}_{chunk_index} — deterministic, so re-indexing overwrites rather "
     "than duplicating."),
    ("Bulk writes. ", "elasticsearch.helpers.bulk() batches an entire corpus into one request."),
    ("Replicas. ", "Set to 0 for a single-node local cluster."),
])

h(doc, "4.5  Retrieval — hybrid search with manual RRF", 2)

body(doc,
     "Elasticsearch's native RRF retriever requires a Platinum licence and throws "
     "AuthorizationException on the basic licence, so fusion is implemented in Python over two "
     "separate queries.")

code(doc,
     "RRF(doc) = 1/(k + rank_bm25 + 1) + 1/(k + rank_knn + 1),   k = 60\n"
     "\n"
     "window          = top_k × 10 = 50 candidates from each retriever\n"
     "kNN quality gate: ES score ≥ 0.65  ⇔  cosine ≥ 0.30\n"
     "relevance gate  : final RRF score ≥ 0.020\n"
     "theoretical max : 2/(60+1) = 0.0328")

bullets(doc, [
    ("Union, not intersection. ", "A document found by only one retriever still receives a partial "
     "score, so a purely lexical or purely semantic match is not discarded."),
    ("Two gates, two purposes. ", "The kNN gate drops individually weak vector hits before fusion. "
     "The RRF gate then drops anything that no retriever ranked well. Together they let the system "
     "return nothing — a query about chocolate cake yields zero results rather than the five "
     "least-bad legal documents."),
    ("Metadata filters. ", "doc_type, date_from and date_to translate to Elasticsearch filter "
     "clauses applied identically to both branches."),
])

h(doc, "4.6  Confidence scoring", 2)

body(doc,
     "The RRF score cannot be shown to a user. Its ceiling is 2/61 ≈ 0.033, so a perfect hit "
     "rendered as a percentage reads as 3% confident. It is also purely ordinal: it records that "
     "one chunk outranked another and nothing about how well either matched. The similarity "
     "Elasticsearch computed was being discarded during fusion.")

body(doc, "retrieval/confidence.py reconstructs a genuine 0–1 score from two signals:")

table(doc, ["Signal", "Definition", "Calibration"], [
    ("Semantic", "Cosine similarity, recovered as 2 × es_score − 1 because dot_product scoring "
     "returns (1 + cosine) / 2.",
     "0.30 → 0%, 0.80 → 100%. The floor is not arbitrary: it is the same threshold the kNN "
     "quality gate already treats as noise."),
    ("Lexical", "BM25 rank, not raw score — BM25 scores are unbounded and not comparable across "
     "queries.",
     "Reciprocal decay halving every 2 ranks, the same shape RRF itself uses."),
], widths=[0.9, 3.0, 2.8])

body(doc,
     "The two combine as a noisy-OR rather than a weighted sum, because they are independent kinds "
     "of evidence and either alone can be conclusive:")

code(doc,
     "semantic = 0.85 × semantic_signal      # reliability of a perfect semantic match alone\n"
     "lexical  = 0.55 × lexical_signal       # reliability of a top BM25 hit alone\n"
     "\n"
     "confidence = (1 − (1 − semantic) × (1 − lexical)) / normaliser\n"
     "\n"
     "bands: ≥ 75% high (green)   ≥ 50% moderate (amber)   < 50% low (red)")

callout(doc, "Why noisy-OR, not a weighted sum:",
        "A weighted sum was implemented first and failed on a real query. Asked what was scheduled "
        "at 3 Allen Center, the correct email had cosine 0.31 at BM25 rank 1, while an off-topic "
        "Super Bowl email had cosine 0.42 at rank 4. Under linear rank decay, rank 4 was worth 94% "
        "of rank 1, so the semantic weight scored the wrong document higher — 49% against 37%. "
        "Proper nouns are carried entirely by BM25; paraphrased questions entirely by the "
        "embedding. Noisy-OR lets either signal carry a result alone, and corroboration falls out "
        "of the combination without a separate bonus term. The case is now a regression test.")

h(doc, "4.7  Document grouping", 2)

body(doc,
     "Retrieval scores chunks, but a 24-chunk opinion frequently contributes three of the top five "
     "hits, which surfaces as the same file cited three times. retrieval/grouping.py merges hits "
     "sharing a source document — by filename for PDFs and emails, by originating filing for LEDGAR "
     "provisions.")

bullets(doc, [
    ("Grouping precedes synthesis. ", "This is the load-bearing detail. If duplicates were merged "
     "only for display, the answer's [3] marker would point at a citation that no longer exists."),
    ("Nothing is discarded. ", "All excerpts still reach the LLM under the group's single citation "
     "number; the synthesiser is told that one source may hold several excerpts separated by [...]."),
    ("Best-chunk metadata wins. ", "Results arrive best-first, so the group inherits the top "
     "chunk's page — the page the viewer should open at. Group confidence is the maximum across "
     "chunks: a document is as good as its strongest match, and averaging would penalise a "
     "document for also matching weakly elsewhere."),
])

h(doc, "4.8  Natural-language interface", 2)

body(doc, "Query understanding — nli/query_processor.py")
body(doc,
     "Claude converts the question into JSON: an intent (factual, conceptual, or filter), a clean "
     "reformulated retrieval string with filter conditions stripped out, and structured filters. "
     "Any API or parse failure falls back to using the raw query, so the NLI layer can never take "
     "search down.")

body(doc, "Answer synthesis — nli/synthesizer.py")
body(doc,
     "Claude writes an answer under 200 words from the retrieved excerpts only, citing sources as "
     "[1], [2]. The prompt requires it to state plainly when the excerpts are insufficient rather "
     "than inferring beyond them. On any error it degrades to showing raw excerpt snippets.")

h(doc, "4.9  Interface and evidence viewer", 2)

bullets(doc, [
    ("Answer first. ", "The cited synthesis appears above the source list; a collapsible panel "
     "shows the parsed intent, the reformulated query and any active filters."),
    ("Confidence ring. ", "An inline SVG donut sits at the right edge of each citation row, using "
     "currentColor at 15% opacity for the track so it renders correctly in light and dark themes. "
     "Each expander also carries a plain-language breakdown of where its score came from."),
    ("In-app PDF viewer. ", "PyMuPDF renders every page to PNG and highlights each matched excerpt "
     "at the word level, with punctuation and case stripped — Docling's OCR output differs "
     "cosmetically from a PDF's own text layer (smart quotes, spacing), so exact matching fails. "
     "Every matched page is highlighted and flagged, not just the first."),
])

doc.add_page_break()

# ── 5. Evaluation ─────────────────────────────────────────────────────────────

h(doc, "5.  Evaluation and Testing", 1)

h(doc, "5.1  Test suite", 2)

table(doc, ["Suite", "Tests", "Covers"], [
    ("test_chunking.py", "36", "All three chunkers in isolation: chunk counts, token budget never "
     "exceeded, table detection and atomicity, metadata propagation, fragment splitting, greedy "
     "merge, KV split, word-level fallback, empty bodies."),
    ("test_retrieval.py", "22", "RRF formula correctness, score ordering, both relevance gates, "
     "top_k capping, empty-list handling, doc_type filter field mapping, merge speed."),
    ("test_confidence.py", "14", "Cosine recovery from ES scores, 0–1 bounds across every signal "
     "combination, monotonicity in each signal, corroboration bonus, single-retriever caps, and "
     "the 3 Allen Center rank-inversion regression."),
    ("test_grouping.py", "10", "Collapse by document, relevance-order preservation, best-chunk "
     "metadata precedence, provisions grouped by filing, no excerpt text lost, inputs not mutated."),
], widths=[1.3, 0.6, 4.8])

body(doc, "Total: 82 tests, all passing.")

body(doc,
     "Unit tests replace Elasticsearch with a MagicMock returning canned hit lists — no network, no "
     "embedding model, fully deterministic. Integration tests are marked separately and run against "
     "a live index.")

h(doc, "5.2  Measured retrieval quality", 2)

body(doc, "Against the live index (956 chunks), after the confidence model was corrected:")

table(doc, ["Query", "Top result", "Cosine", "BM25", "Confidence"], [
    ("foreign judgment recognition property in Washington", "alterna_v_spicejet.pdf", "0.76", "#2", "94%"),
    ("Arotin summary judgment standing statute of frauds", "arotin_v._arotin.pdf", "0.68", "#2", "86%"),
    ("EWEB Schedule Master Agreement", "email_3.txt", "0.48", "#1", "74%"),
    ("scheduled event 3 Allen Center weekend", "email_2.txt (correct)", "0.31", "#1", "60%"),
    ("  — same query, off-topic result", "email_4.txt (Super Bowl)", "0.42", "#4", "41%"),
    ("governing law clause New York", "LEDGAR provisions", "0.49–0.56", "#3–#7", "47–58%"),
    ("recipe for chocolate cake", "— no results —", "—", "—", "gated"),
], widths=[2.2, 1.5, 0.7, 0.5, 0.8])

body(doc,
     "The pattern is the intended one. Precise legal questions with strong lexical and semantic "
     "agreement score high; generic boilerplate scores moderate, honestly reflecting that hundreds "
     "of governing-law clauses are near-identical; off-topic queries return nothing at all.")

h(doc, "5.3  Performance", 2)

table(doc, ["Measure", "Target", "Result"], [
    ("End-to-end search (embed + BM25 + kNN + RRF)", "< 2 s", "Passing"),
    ("RRF merge over 50 candidates (Python)", "< 50 ms", "Passing"),
    ("Full unit suite", "—", "~10 s"),
    ("Re-ingestion of unchanged PDFs", "—", "OCR skipped via hash match"),
], widths=[3.0, 1.2, 2.5])

h(doc, "5.4  Defects found and fixed during review", 2)

table(doc, ["Defect", "Root cause", "Resolution"], [
    ("Compliance citations displayed as \"Unknown\"",
     "JsonChunk carried no provenance fields; the source filing existed only as text inside the "
     "chunk body.",
     "Source filing and clause labels captured at ingest; UI falls back to parsing them out of "
     "chunk text so the ~823 already-indexed chunks display correctly without re-ingestion."),
    ("Same document cited three times",
     "Retrieval scores chunks, not documents.",
     "Grouping by source document, applied before synthesis so citation numbers stay valid."),
    ("Confidence shown as ~3% for correct documents",
     "The RRF score was being rendered as a percentage; its ceiling is 0.0328.",
     "Confidence reconstructed from cosine similarity and BM25 rank."),
    ("Confidence ranked the wrong document higher",
     "Linear rank decay made BM25 rank 4 worth 94% of rank 1, letting semantic weight override a "
     "proper-noun match.",
     "Reciprocal decay plus noisy-OR combination; locked by a regression test."),
    ("Email searches returned nothing at all",
     "Only the PDF chunker sets doc_type. Filtering emails on doc_type matched zero documents and, "
     "as a filter clause, emptied both retrieval branches — indistinguishable from nothing being "
     "indexed.",
     "Email filters on chunk_type, covering both email and email_fragment. A test now asserts each "
     "filter's target field exists on the dataclass it selects."),
], widths=[1.6, 2.3, 2.8])

doc.add_page_break()

# ── 6. Output ─────────────────────────────────────────────────────────────────

h(doc, "6.  Output — A Worked Example", 1)

body(doc, "Query: \"What were the grounds for summary judgment in the Arotin case?\"")

body(doc, "Step 1 — Query understanding")
code(doc,
     '{"intent": "factual",\n'
     ' "reformulated_query": "grounds summary judgment Arotin",\n'
     ' "filters": {"doc_type": null, "date_from": null, "date_to": null,\n'
     '             "party_names": ["Arotin"]},\n'
     ' "explanation": "Seeking the specific legal grounds on which summary\n'
     '                 judgment was granted in the Arotin case."}')

body(doc, "Step 2 — Retrieval and fusion")
body(doc,
     "BM25 and kNN each return 50 candidates. RRF fuses them; four chunks clear the 0.020 gate. "
     "Two are from arotin_v._arotin.pdf (pages 1 and 3) and are collapsed into a single cited "
     "source before synthesis.")

body(doc, "Step 3 — Synthesised answer")
code(doc,
     "The trial court denied the plaintiffs-appellants' (William and Linda Arotin)\n"
     "motion for summary judgment and granted summary judgment to the\n"
     "defendants-appellees (Randy J. Arotin and Faith O. Cull-Arotin) [1].\n"
     "\n"
     "The appellees' grounds centered on lack of standing, arguing that appellants\n"
     "had not pleaded facts establishing a legally enforceable lien or mortgage\n"
     "necessary to bring a foreclosure action [1]. Additional grounds included:\n"
     "\n"
     "  • the alleged oral contract was barred by the statute of frauds [1];\n"
     "  • the quitclaim deed barred parol evidence of consideration beyond $10.00 [1];\n"
     "  • the county auditor's exemption form suggested a gift, not a sale [1].\n"
     "\n"
     "The Court of Appeals of Ohio, Eleventh Appellate District, affirmed on\n"
     "April 6, 2026 [1].")

body(doc, "Step 4 — Source list")
code(doc,
     "Sources  (3 documents)\n"
     "\n"
     "[1] arotin_v._arotin.pdf · p.1, 3 · IN THE COURT OF APPEALS OF OHIO      ( 86% )\n"
     "    ELEVENTH APPELLATE DIST · April 6, 2026 — 86% confidence  section\n"
     "    ├─ top excerpt shown inline\n"
     "    ├─ \"+1 more matching excerpt on p.3 — open it to read them in context.\"\n"
     "    └─ 86% confidence (high) — semantic similarity 0.68 · keyword rank #2\n"
     "                              · matched by both retrievers · RRF 0.032258\n"
     "\n"
     "[2] 2019/QTR1/000119312519044328/d691151dex101.htm · releases         ( 53% )\n"
     "    — 53% confidence  json_provision\n"
     "\n"
     "[3] alterna_aircraft_v_b_ltd._v._spicejet_ltd..pdf · p.9 ·            ( 41% )\n"
     "    April 9, 2026 — 41% confidence  section")

body(doc,
     "Note that the citation numbering in the answer matches the source list exactly, and that the "
     "two Arotin excerpts share one number. Clicking through opens the PDF with both page 1 and "
     "page 3 highlighted at the matched passages.")

body(doc,
     "Confidence is not monotonic with list position — the list is ordered by RRF, while the "
     "percentages measure match strength, and the two can legitimately disagree. Sorting the list "
     "by confidence would demote correct-but-lexically-found documents, which is precisely the "
     "failure mode the noisy-OR model exists to prevent.",
     italic=True)

doc.add_page_break()

# ── 7. Limitations ────────────────────────────────────────────────────────────

h(doc, "7.  Known Limitations and Next Steps", 1)

h(doc, "7.1  Open defects", 2)

table(doc, ["Issue", "Detail", "Fix"], [
    ("Date filters do not work", "PDF dates are stored as prose (\"April 10, 2026\") in a keyword "
     "field, so range queries compare lexicographically and return meaningless results. Emails "
     "store their date under date, not document_date, so any date filter silently excludes every "
     "email. 51 of 126 PDF chunks — all the scanned ones — have no date at all.",
     "Normalise to ISO 8601 at ingest, map as date, and filter both fields."),
    ("Committed API key", "A live-looking Anthropic API key is present in .env.example, which is "
     "tracked by git.",
     "Rotate the key immediately and replace the file's value with a placeholder."),
    ("party_names extracted but unused", "The NLI layer returns party names; the filter builder "
     "ignores them. They are only appended to the retrieval string.",
     "Either apply as a filter or remove from the schema."),
    ("ingestion/json_ingester.py is empty", "A placeholder module; JSON ingestion happens inside "
     "the chunker.", "Delete or implement for symmetry with the other two ingesters."),
    ("Config drift", "config/settings.py declares ES_INDEX = \"legal_docs\" and "
     "EMBEDDING_MODEL, but storage/es_client.py hardcodes \"legallens\" and the embedder hardcodes "
     "the model name. The settings module is not read.",
     "Route both through config, or delete the unused module."),
], widths=[1.5, 3.3, 1.9])

h(doc, "7.2  Design limitations", 2)

bullets(doc, [
    ("Relevance gate is corpus-calibrated. ", "MIN_RRF_SCORE = 0.020 suits the current corpus. RRF "
     "scores compress as a corpus grows, so this will need lowering at scale."),
    ("Confidence weights are judgement calls. ", "The reliability constants are not a calibrated "
     "probability model. Calibrating them against human relevance judgements would make the "
     "percentages defensible as probabilities rather than as a ranking aid."),
    ("Embeddings are general-purpose. ", "all-mpnet-base-v2 is not fine-tuned on legal text, which "
     "shows in the 3 Allen Center case, where the correct document scored barely above the noise "
     "floor. A legal-domain model would raise the semantic floor across the corpus."),
    ("Cosine ceiling is an estimate. ", "COSINE_CEIL = 0.80 was set from the model's general "
     "behaviour, not measured against this corpus."),
    ("No access control enforcement. ", "The schema is designed for it — matter_id and "
     "clearance_level fit the existing mapping without re-indexing — but no authentication layer "
     "exists."),
    ("Only 823 of 374,639 LEDGAR chunks are indexed. ", "Full ingestion is an embedding-throughput "
     "exercise, not an architectural one."),
])

h(doc, "7.3  Recommended next steps", 2)

bullets(doc, [
    "Rotate the committed API key and replace it with a placeholder.",
    "Normalise dates to ISO 8601 at ingest so date filtering works.",
    "Add doc_type to EmailChunk and JsonChunk so all three filters use one uniform mechanism.",
    "Build a labelled relevance set from the corpus and measure precision@5 and recall — the "
    "current evaluation demonstrates correct behaviour on selected queries but does not quantify "
    "retrieval quality.",
    "Calibrate the confidence constants against those judgements.",
    "Add a cross-encoder re-ranking stage over the fused top 50.",
    "Re-enable Elasticsearch security and add the authentication layer the access-control schema "
    "assumes.",
])

doc.add_page_break()

# ── 8. Appendix ───────────────────────────────────────────────────────────────

h(doc, "8.  Appendix — Module Reference", 1)

table(doc, ["Module", "Lines", "Responsibility"], [
    ("app.py", "~250", "Streamlit interface, confidence rings, source list, PDF viewer dialog"),
    ("ingest.py", "131", "Pipeline orchestration for all three corpora"),
    ("search_cli.py", "31", "Interactive terminal search"),
    ("pdf_viewer.py", "~175", "PyMuPDF rendering and multi-page word-level highlighting"),
    ("ingestion/pdf_ingester.py", "220", "Docling OCR, hash-based staleness, page markers, dates"),
    ("ingestion/email_ingester.py", "~200", "Enron-format parser, headers, reply chains, threading"),
    ("chunking/pdf_chunker.py", "387", "RSM section tree, atomic tables, three merge passes"),
    ("chunking/email_chunker.py", "~180", "One email per chunk, paragraph-boundary fragments"),
    ("chunking/json_chunker.py", "~200", "Streaming LEDGAR chunker, citation metadata"),
    ("embedding/embedder.py", "59", "all-mpnet-base-v2 singleton, batched normalised encoding"),
    ("storage/indexer.py", "106", "Index mapping, creation, bulk indexing"),
    ("storage/es_client.py", "9", "Elasticsearch client factory"),
    ("retrieval/search.py", "147", "Hybrid BM25 + kNN, manual RRF, filters, gates"),
    ("retrieval/confidence.py", "101", "Noisy-OR confidence from cosine and BM25 rank"),
    ("retrieval/grouping.py", "82", "Collapse chunk hits into per-document citations"),
    ("nli/query_processor.py", "~110", "Claude query understanding with graceful fallback"),
    ("nli/synthesizer.py", "~100", "Claude cited answer synthesis with graceful fallback"),
    ("tests/", "~700", "82 tests across four suites"),
], widths=[2.0, 0.7, 4.0])

body(doc,
     "Full architectural rationale, including alternatives considered and trade-offs accepted for "
     "each decision, is recorded in DECISIONS.md in the repository.",
     italic=True)

doc.save(OUTPUT)
print(f"Written: {OUTPUT}")
